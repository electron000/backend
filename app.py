import sqlite3
import os
import io
import math
import threading
import logging
import shutil
import uuid
from datetime import datetime, timezone, timedelta
from functools import wraps

from flask import Flask, request, jsonify, send_file, g
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

import jwt
import pandas as pd
import re

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import landscape, A4, A3, A2
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "https://mainsys.vercel.app"}})

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'a-very-secure-fallback-secret-key-for-dev')
DB_FILE = os.environ.get('DATABASE_FILE', 'contracts.db')
EXCEL_FILE = os.environ.get('EXCEL_FILE', 'Contract Details.xlsx')
ADMIN_DEFAULT_PASSWORD = os.environ.get('ADMIN_DEFAULT_PASSWORD')
USER_DEFAULT_PASSWORD = os.environ.get('USER_DEFAULT_PASSWORD')
TABLE_NAME = 'contracts'
UPLOAD_FOLDER = '.'
BACKUP_FOLDER = os.path.join(UPLOAD_FOLDER, 'backups')
MAX_BACKUPS = 5
ALLOWED_EXTENSIONS = {'xlsx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
FIXED_START_COL = 'SL No'

file_lock = threading.Lock()
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def create_backup():
    if not os.path.exists(BACKUP_FOLDER):
        os.makedirs(BACKUP_FOLDER)
    
    source_file = os.path.join(UPLOAD_FOLDER, EXCEL_FILE)
    
    if os.path.exists(source_file):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_file_name = f"backup_{timestamp}_{EXCEL_FILE}"
        destination = os.path.join(BACKUP_FOLDER, backup_file_name)
        
        try:
            shutil.copy(source_file, destination)
            logging.info(f"Created backup: {destination}")
            
            backups = sorted(
                [os.path.join(BACKUP_FOLDER, f) for f in os.listdir(BACKUP_FOLDER)],
                key=os.path.getmtime
            )
            if len(backups) > MAX_BACKUPS:
                os.remove(backups[0])
                logging.info(f"Removed old backup: {backups[0]}")

        except Exception as e:
            logging.error(f"Could not create backup. Error: {e}")

def get_db_connection():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn

def get_sqlite_column_type_string(data_type):
    return 'REAL' if data_type == 'numeric' else 'TEXT'

def sanitize_column_name(col_name):
    return ''.join(e for e in col_name if e.isalnum() or e.isspace()).strip().replace(' ', '_')

def get_current_schema():
    conn = get_db_connection()
    try:
        schema_rows = conn.execute("SELECT column_uuid, name, type FROM schema_metadata ORDER BY display_order").fetchall()
    except sqlite3.OperationalError:
        return [], {'numeric': [], 'date': [], 'text': []}, []
    finally:
        conn.close()
    headers = [row['name'] for row in schema_rows]
    field_types = {'numeric': [], 'date': [], 'text': []}
    for row in schema_rows:
        field_types.setdefault(row['type'], []).append(row['name'])
    return headers, field_types, schema_rows

def infer_column_types(df):
    inferred = {"numeric": [], "date": [], "text": []}
    date_pattern = re.compile(r'^\d{1,4}[-/]\d{1,2}[-/]\d{1,4}$')

    for col in df.columns:
        col_data = df[col].dropna().astype(str)

        if col_data.empty:
            inferred["text"].append(col)
            continue
        if any(keyword in col.lower() for keyword in ['(₹)', 'amount', 'value', 'duration', 'charges', '(yr)', 'years', 'year', '(inr)']):
            inferred["numeric"].append(col)
            continue
        if 'date' in col.lower():
            inferred["date"].append(col)
            continue
        sample_size = min(100, len(col_data))
        sample = col_data.sample(sample_size)
        numeric_count = 0
        date_count = 0
        for item in sample:
            try:
                float(item)
                numeric_count += 1
                continue
            except (ValueError, TypeError):
                pass
            if date_pattern.match(item):
                try:
                    pd.to_datetime(item, errors='raise')
                    date_count += 1
                except (ValueError, TypeError):
                    pass   
        if (numeric_count / sample_size) >= 0.9:
            inferred["numeric"].append(col)
        elif (date_count / sample_size) >= 0.9:
            inferred["date"].append(col)
        else:
            inferred["text"].append(col)           
    return inferred

def reindex_sl_no_in_db(cursor):
    sanitized_sl_no = sanitize_column_name(FIXED_START_COL)
    try:
        rows_to_reindex = cursor.execute(f"SELECT rowid FROM {TABLE_NAME} ORDER BY rowid").fetchall()
        for index, row in enumerate(rows_to_reindex):
            cursor.execute(f'UPDATE {TABLE_NAME} SET "{sanitized_sl_no}" = ? WHERE rowid = ?', (str(index + 1), row['rowid']))
    except sqlite3.OperationalError as e:
        logging.warning(f"Could not re-index '{FIXED_START_COL}'. Error: {e}")

def initialize_schema_from_excel(cursor):
    df = pd.read_excel(EXCEL_FILE, dtype=str).fillna('')
    if FIXED_START_COL not in df.columns:
        df.insert(0, FIXED_START_COL, "")
    inferred_types = infer_column_types(df)
    cursor.execute("DELETE FROM schema_metadata")
    for i, col_name in enumerate(df.columns):
        col_type = next((t for t, names in inferred_types.items() if col_name in names), 'text')
        cursor.execute("INSERT INTO schema_metadata (column_uuid, name, type, display_order) VALUES (?, ?, ?, ?)", (str(uuid.uuid4()), col_name, col_type, i))
    schema_rows = cursor.execute("SELECT name, type FROM schema_metadata ORDER BY display_order").fetchall()
    create_cols_sql = [f'"{sanitize_column_name(row["name"])}" {get_sqlite_column_type_string(row["type"])}' for row in schema_rows]
    cursor.execute(f"DROP TABLE IF EXISTS {TABLE_NAME}")
    cursor.execute(f"CREATE TABLE {TABLE_NAME} ({', '.join(create_cols_sql)})")
    df_sanitized_cols = df.copy()
    df_sanitized_cols.columns = [sanitize_column_name(col) for col in df.columns]
    df_sanitized_cols.to_sql(TABLE_NAME, cursor.connection, if_exists='append', index=False)
    reindex_sl_no_in_db(cursor)

def setup_database_and_schema():
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='users'")
        if cursor.fetchone() is None:
            cursor.execute('CREATE TABLE users (id INTEGER PRIMARY KEY, username TEXT UNIQUE NOT NULL, password TEXT NOT NULL, role TEXT NOT NULL)')
            hashed_admin_pass = generate_password_hash(ADMIN_DEFAULT_PASSWORD)
            cursor.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)", ('Infocom-Admin', hashed_admin_pass, 'admin'))
            hashed_user_pass = generate_password_hash(USER_DEFAULT_PASSWORD)
            cursor.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)", ('Infocom-User', hashed_user_pass, 'user'))
            logging.info("Users table created and default users added.")
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='schema_metadata'")
        if cursor.fetchone() is None:
            cursor.execute('CREATE TABLE schema_metadata (id INTEGER PRIMARY KEY, column_uuid TEXT UNIQUE NOT NULL, name TEXT UNIQUE NOT NULL, type TEXT NOT NULL, display_order INTEGER NOT NULL)')
            logging.info("Schema metadata table created with column_uuid.")
        else:
            cursor.execute("PRAGMA table_info(schema_metadata)")
            columns = [col[1] for col in cursor.fetchall()]
            if 'column_uuid' not in columns:
                cursor.execute("ALTER TABLE schema_metadata ADD COLUMN column_uuid TEXT")
                existing_schemas = cursor.execute("SELECT id FROM schema_metadata WHERE column_uuid IS NULL").fetchall()
                for row in existing_schemas:
                    cursor.execute("UPDATE schema_metadata SET column_uuid = ? WHERE id = ?", (str(uuid.uuid4()), row['id']))
                cursor.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_schema_metadata_column_uuid ON schema_metadata (column_uuid)")
                logging.info("Added and populated column_uuid in schema_metadata.")
        cursor.execute(f"SELECT name FROM sqlite_master WHERE type='table' AND name='{TABLE_NAME}'")
        if cursor.fetchone() is None:
            if os.path.exists(EXCEL_FILE):
                initialize_schema_from_excel(cursor)
            else:
                default_cols = [(FIXED_START_COL, 'numeric'), ('Contract Name', 'text'), ('Value', 'numeric'), ('Start Date', 'date')]
                for i, (name, type) in enumerate(default_cols):
                    cursor.execute("INSERT INTO schema_metadata (column_uuid, name, type, display_order) VALUES (?, ?, ?, ?)", (str(uuid.uuid4()), name, type, i))
                schema = cursor.execute("SELECT name, type FROM schema_metadata ORDER BY display_order").fetchall()
                sanitized_cols_with_types = [f'"{sanitize_column_name(col["name"])}" {get_sqlite_column_type_string(col["type"])}' for col in schema]
                cursor.execute(f"CREATE TABLE {TABLE_NAME} ({', '.join(sanitized_cols_with_types)})")

def export_db_to_excel():
    pass

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('authorization', '').split(" ")[-1] if 'authorization' in request.headers else None
        if not token: return jsonify({'message': 'Token is missing!'}), 401
        try:
            data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=["HS256"])
            g.current_user, g.current_role = data['username'], data['role']
        except (jwt.ExpiredSignatureError, jwt.InvalidTokenError, jwt.PyJWTError) as e:
            logging.warning(f"JWT validation failed: {e}")
            return jsonify({'message': 'Token is invalid or expired!'}), 401
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('authorization', '').split(" ")[-1] if 'authorization' in request.headers else None
        if not token: return jsonify({'message': 'Token is missing!'}), 401
        try:
            data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=["HS256"])
            if data.get('role') != 'admin':
                return jsonify({'message': 'Admin privileges required!'}), 403
            g.current_user, g.current_role = data['username'], data['role']
        except (jwt.ExpiredSignatureError, jwt.InvalidTokenError, jwt.PyJWTError) as e:
            logging.warning(f"JWT validation failed: {e}")
            return jsonify({'message': 'Token is invalid or expired!'}), 401
        return f(*args, **kwargs)
    return decorated

@app.route('/api/schema', methods=['GET'])
@admin_required
def get_schema_endpoint():
    _, _, schema_rows_with_uuid = get_current_schema() 
    schema = [{'column_uuid': row['column_uuid'], 'name': row['name'], 'type': row['type']} for row in schema_rows_with_uuid]
    return jsonify(schema)

@app.route('/api/schema/columns', methods=['POST'])
@admin_required
def add_column():
    data = request.get_json()
    new_col_name = data.get('name')
    new_col_type = data.get('type')
    if not new_col_name or not new_col_type:
        return jsonify({"error": "New column name and type are required."}), 400
    
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM schema_metadata WHERE LOWER(name) = LOWER(?)", (new_col_name,))
            if cursor.fetchone():
                return jsonify({"error": f"A column named '{new_col_name}' already exists."}), 409
            
            new_column_uuid = str(uuid.uuid4()) 
            sqlite_type_str = get_sqlite_column_type_string(new_col_type)
            
            cursor.execute(f'ALTER TABLE {TABLE_NAME} ADD COLUMN "{sanitize_column_name(new_col_name)}" {sqlite_type_str}')
            max_order_result = cursor.execute("SELECT MAX(display_order) FROM schema_metadata").fetchone()
            max_order = max_order_result[0] if max_order_result and max_order_result[0] is not None else -1
            
            cursor.execute("INSERT INTO schema_metadata (column_uuid, name, type, display_order) VALUES (?, ?, ?, ?)", (new_column_uuid, new_col_name, new_col_type, max_order + 1))
            logging.info(f"Column '{new_col_name}' added successfully.")
    except sqlite3.Error as e:
        logging.error(f"Database error adding column '{new_col_name}': {e}")
        return jsonify({"error": f"Database error: {e}"}), 500
    
    export_db_to_excel()
    return jsonify({"message": f"Column '{new_col_name}' added successfully.", "new_uuid": new_column_uuid}), 201

@app.route('/api/schema/columns/<string:column_uuid>', methods=['PUT'])
@admin_required
def update_column(column_uuid):
    data = request.get_json()
    new_name = data.get('name')
    new_type = data.get('type')

    if not new_name and not new_type:
        return jsonify({"error": "No new name or type provided for update."}), 400

    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            
            cursor.execute("SELECT name, type FROM schema_metadata WHERE column_uuid = ?", (column_uuid,))
            current_col_metadata = cursor.fetchone()
            if not current_col_metadata:
                return jsonify({"error": f"Column with UUID '{column_uuid}' not found."}), 404
            
            column_name = current_col_metadata['name']

            if new_name and new_name != column_name:
                cursor.execute("SELECT name FROM schema_metadata WHERE LOWER(name) = LOWER(?) AND column_uuid != ?", (new_name, column_uuid))
                if cursor.fetchone():
                    return jsonify({"error": f"A column named '{new_name}' already exists."}), 409

            current_schema_rows = cursor.execute("SELECT column_uuid, name, type FROM schema_metadata ORDER BY display_order").fetchall()
            
            target_schema = []
            for col in current_schema_rows:
                if col['column_uuid'] == column_uuid:
                    target_schema.append({'column_uuid': col['column_uuid'],'name': new_name if new_name else col['name'],'type': new_type if new_type else col['type']})
                else:
                    target_schema.append(dict(col))

            temp_table_name = f"{TABLE_NAME}_temp_update"
            new_col_definitions = [f'"{sanitize_column_name(c["name"])}" {get_sqlite_column_type_string(c["type"])}' for c in target_schema]
            cursor.execute(f"CREATE TABLE {temp_table_name} ({', '.join(new_col_definitions)})")

            old_sanitized_names = [f'"{sanitize_column_name(c["name"])}"' for c in current_schema_rows]
            new_sanitized_names = [f'"{sanitize_column_name(c["name"])}"' for c in target_schema]
            cursor.execute(f"INSERT INTO {temp_table_name} ({', '.join(new_sanitized_names)}) SELECT {', '.join(old_sanitized_names)} FROM {TABLE_NAME}")

            cursor.execute(f"DROP TABLE {TABLE_NAME}")
            cursor.execute(f"ALTER TABLE {temp_table_name} RENAME TO {TABLE_NAME}")

            if new_name and new_name != column_name:
                cursor.execute("UPDATE schema_metadata SET name = ? WHERE column_uuid = ?", (new_name, column_uuid))
            if new_type:
                cursor.execute("UPDATE schema_metadata SET type = ? WHERE column_uuid = ?", (new_type, column_uuid))
            
            logging.info(f"Column '{column_name}' updated successfully.")

    except sqlite3.Error as e:
        logging.error(f"Error updating column (UUID: {column_uuid}): {e}")
        return jsonify({"error": str(e)}), 500
        
    export_db_to_excel()
    return jsonify({"message": "Column updated successfully."}), 200

@app.route('/api/schema/columns/<string:column_uuid_to_delete>', methods=['DELETE'])
@admin_required
def delete_column(column_uuid_to_delete):
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM schema_metadata WHERE column_uuid = ?", (column_uuid_to_delete,))
            column_to_delete_info = cursor.fetchone()
            if not column_to_delete_info:
                return jsonify({"error": f"Column with UUID '{column_uuid_to_delete}' not found."}), 404

            column_name_to_delete = column_to_delete_info['name']
            if sanitize_column_name(column_name_to_delete) == sanitize_column_name(FIXED_START_COL):
                return jsonify({"error": "Cannot delete the fixed 'SL No' column."}), 403

            remaining_schema_rows = cursor.execute("SELECT column_uuid, name, type FROM schema_metadata WHERE column_uuid != ? ORDER BY display_order", (column_uuid_to_delete,)).fetchall()

            if not remaining_schema_rows:
                return jsonify({"error": "Cannot delete the last remaining column."}), 400

            temp_table_name = f"{TABLE_NAME}_temp_delete"
            new_col_definitions = [f'"{sanitize_column_name(r["name"])}" {get_sqlite_column_type_string(r["type"])}' for r in remaining_schema_rows]
            cursor.execute(f"CREATE TABLE {temp_table_name} ({', '.join(new_col_definitions)})")

            cols_to_copy = [f'"{sanitize_column_name(r["name"])}"' for r in remaining_schema_rows]
            cursor.execute(f"INSERT INTO {temp_table_name} ({', '.join(cols_to_copy)}) SELECT {', '.join(cols_to_copy)} FROM {TABLE_NAME}")

            cursor.execute(f"DROP TABLE {TABLE_NAME}")
            cursor.execute(f"ALTER TABLE {temp_table_name} RENAME TO {TABLE_NAME}")

            cursor.execute("DELETE FROM schema_metadata WHERE column_uuid = ?", (column_uuid_to_delete,))
            
            logging.info(f"Column '{column_name_to_delete}' deleted successfully.")

    except sqlite3.Error as e:
        logging.error(f"Error deleting column (UUID: {column_uuid_to_delete}): {e}")
        return jsonify({"error": str(e)}), 500
        
    export_db_to_excel()
    return jsonify({"message": f"Column '{column_name_to_delete}' deleted successfully."}), 200

@app.route('/api/schema/reorder', methods=['POST'])
@admin_required
def reorder_columns():
    ordered_column_uuids = request.get_json()
    if not ordered_column_uuids:
        return jsonify({"error": "A list of ordered column UUIDs is required."}), 400
    
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            for i, col_uuid in enumerate(ordered_column_uuids):
                cursor.execute("UPDATE schema_metadata SET display_order = ? WHERE column_uuid = ?", (i, col_uuid))
            logging.info("Column order updated successfully.")
    except sqlite3.Error as e:
        logging.error(f"Error reordering columns: {e}")
        return jsonify({"error": str(e)}), 500

    export_db_to_excel()
    return jsonify({"message": "Column order updated successfully."}), 200

@app.route('/api/login', methods=['POST'])
def login():
    data = request.get_json()
    username, password, login_mode = data.get('username'), data.get('password'), data.get('loginMode')
    if not all([username, password, login_mode]): return jsonify({"error": "Missing credentials or login mode."}), 400
    with get_db_connection() as conn:
        user = conn.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
    if user and check_password_hash(user['password'], password):
        if user['role'] != login_mode: return jsonify({"error": f"Please use the '{user['role'].capitalize()}' login panel."}), 403
        token = jwt.encode({'username': user['username'], 'role': user['role'], 'exp': datetime.now(timezone.utc) + timedelta(hours=24)}, app.config['SECRET_KEY'], algorithm="HS256")
        return jsonify({'token': token}), 200
    return jsonify({"error": "Invalid credentials"}), 401

@app.route('/api/contracts', methods=['GET'])
@token_required 
def get_contracts():
    headers, field_types, _ = get_current_schema()
    if not headers:
        return jsonify({"data": [], "totalPages": 0, "currentPage": 1, "headers": [], "fieldTypes": {}})
    params = request.args
    where_clauses, query_params = [], []
    filter_field = params.get('filterField')
    if filter_field:
        sanitized_filter_field = sanitize_column_name(filter_field)
        if filter_field in field_types.get('numeric', []):
            min_range, max_range = params.get('minRange'), params.get('maxRange')
            if min_range: where_clauses.append(f'CAST("{sanitized_filter_field}" AS REAL) >= ?'); query_params.append(float(min_range))
            if max_range: where_clauses.append(f'CAST("{sanitized_filter_field}" AS REAL) <= ?'); query_params.append(float(max_range))
        elif filter_field in field_types.get('date', []):
            from_date, to_date = params.get('fromDate'), params.get('toDate')
            if from_date: where_clauses.append(f'date("{sanitized_filter_field}") >= date(?)'); query_params.append(from_date)
            if to_date: where_clauses.append(f'date("{sanitized_filter_field}") <= date(?)'); query_params.append(to_date)
        else:
            filter_value = params.get('filterValue')
            if filter_value: where_clauses.append(f'LOWER("{sanitized_filter_field}") LIKE ?'); query_params.append(f"%{filter_value.lower()}%")
    where_statement = "WHERE " + " AND ".join(where_clauses) if where_clauses else ""
    sort_field = params.get('sortField', FIXED_START_COL)
    sort_direction = params.get('sortDirection', 'asc').upper()
    sanitized_sort_field = sanitize_column_name(sort_field)
    order_by_clause = f'ORDER BY CAST("{sanitize_column_name(FIXED_START_COL)}" AS REAL) ASC'
    if sort_field in field_types.get('numeric', []): order_by_clause = f'ORDER BY CAST("{sanitized_sort_field}" AS REAL) {sort_direction}'
    elif sort_field in field_types.get('date', []): order_by_clause = f'ORDER BY date("{sanitized_sort_field}") {sort_direction}'
    elif sort_field in field_types.get('text', []): order_by_clause = f'ORDER BY "{sanitized_sort_field}" {sort_direction}'
    sanitized_headers = [sanitize_column_name(h) for h in headers]
    select_clause = "SELECT rowid as id, " + ", ".join(f'"{col}"' for col in sanitized_headers)
    with get_db_connection() as conn:
        cursor = conn.cursor()
        count_query = f"SELECT COUNT(*) FROM {TABLE_NAME} {where_statement}"
        total_records = cursor.execute(count_query, tuple(query_params)).fetchone()[0]
        if 'page' not in params:
            data_query = f"{select_clause} FROM {TABLE_NAME} {where_statement} {order_by_clause}"
            results = cursor.execute(data_query, tuple(query_params)).fetchall()
        else:
            page, limit = params.get('page', 1, type=int), params.get('limit', 10, type=int)
            offset = (page - 1) * limit
            data_query = f"{select_clause} FROM {TABLE_NAME} {where_statement} {order_by_clause} LIMIT ? OFFSET ?"
            final_params = tuple(query_params) + (limit, offset)
            results = cursor.execute(data_query, final_params).fetchall()
    sanitized_to_original_map = {sanitize_column_name(h): h for h in headers}
    contracts = []
    for row in results:
        row_dict = dict(row)
        formatted_row = {'id': row_dict.get('id')}
        for s_key, value in row_dict.items():
            if s_key != 'id':
                original_key = sanitized_to_original_map.get(s_key)
                if original_key:
                    formatted_row[original_key] = value
        contracts.append(formatted_row)
    if 'page' not in params:
        return jsonify(contracts)
    else:
        page = params.get('page', 1, type=int)
        limit = params.get('limit', 10, type=int)
        total_pages = math.ceil(total_records / limit) if total_records > 0 else 1
        return jsonify({"data": contracts, "totalPages": total_pages, "currentPage": page, "headers": headers, "fieldTypes": field_types})

@app.route('/api/contracts', methods=['POST'])
@admin_required
def add_contract():
    new_data = request.get_json()
    headers, field_types, _ = get_current_schema()
    is_valid, error_message = validate_data(new_data, field_types)
    if not is_valid: return jsonify({"error": error_message}), 400
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            final_data = {h: new_data.get(h, '') for h in headers if h != FIXED_START_COL}
            sanitized_keys = [sanitize_column_name(k) for k in final_data.keys()]
            columns = ', '.join(f'"{k}"' for k in sanitized_keys)
            placeholders = ', '.join(['?'] * len(final_data))
            cursor.execute(f"INSERT INTO {TABLE_NAME} ({columns}) VALUES ({placeholders})", list(final_data.values()))
            reindex_sl_no_in_db(cursor)
    except sqlite3.Error as e:
        logging.error(f"Database error adding contract: {e}")
        return jsonify({"error": f"Database error: {e}"}), 500
    export_db_to_excel()
    return jsonify({"message": "Contract added successfully"}), 201

@app.route('/api/contracts/<int:row_id>', methods=['PUT'])
@admin_required
def update_contract(row_id):
    updated_data = request.get_json()
    _, field_types, _ = get_current_schema()
    is_valid, error_message = validate_data(updated_data, field_types)
    if not is_valid: return jsonify({"error": error_message}), 400
    try:
        with get_db_connection() as conn:
            set_clauses = [f'"{sanitize_column_name(k)}" = ?' for k in updated_data if k != FIXED_START_COL and k != 'id']
            params = [v for k, v in updated_data.items() if k != FIXED_START_COL and k != 'id']
            if not set_clauses: return jsonify({"error": "No valid fields to update"}), 400
            params.append(row_id)
            conn.execute(f"UPDATE {TABLE_NAME} SET {', '.join(set_clauses)} WHERE rowid = ?", tuple(params))
    except sqlite3.Error as e:
        logging.error(f"Database error updating contract {row_id}: {e}")
        return jsonify({"error": f"Database error: {e}"}), 500
    export_db_to_excel()
    return jsonify({"message": f"Contract {row_id} updated successfully."}), 200

@app.route('/api/contracts/<int:row_id>', methods=['DELETE'])
@admin_required
def delete_contract(row_id):
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(f"DELETE FROM {TABLE_NAME} WHERE rowid = ?", (row_id,))
            reindex_sl_no_in_db(cursor)
    except sqlite3.Error as e:
        logging.error(f"Database error deleting contract {row_id}: {e}")
        return jsonify({"error": f"Database error: {e}"}), 500
    export_db_to_excel()
    return jsonify({"message": f"Contract {row_id} deleted successfully."}), 200

# --- THIS IS YOUR REVERT ENDPOINT ---
@app.route('/api/revert', methods=['POST'])
@admin_required
def revert_to_last_backup():
    try:
        if not os.path.exists(BACKUP_FOLDER) or not os.listdir(BACKUP_FOLDER):
            return jsonify({"error": "No backups available to revert to."}), 404

        backups = [os.path.join(BACKUP_FOLDER, f) for f in os.listdir(BACKUP_FOLDER)]
        latest_backup = max(backups, key=os.path.getctime)

        with file_lock:
            destination_file = os.path.join(UPLOAD_FOLDER, EXCEL_FILE)
            shutil.copy(latest_backup, destination_file)
            logging.info(f"Reverted to backup: {latest_backup}")

            with get_db_connection() as conn:
                cursor = conn.cursor()
                initialize_schema_from_excel(cursor)
            logging.info("Database re-initialized from reverted file.")

        return jsonify({"message": f"Successfully reverted to the backup from {datetime.fromtimestamp(os.path.getctime(latest_backup)).strftime('%Y-%m-%d %H:%M:%S')}"}), 200

    except Exception as e:
        logging.error(f"Error during revert operation: {e}")
        return jsonify({"error": str(e)}), 500

# --- THIS IS YOUR UPDATED UPLOAD ENDPOINT ---
@app.route('/api/upload', methods=['POST'])
@admin_required
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({"error": "Invalid or no selected file"}), 400
    try:
        with file_lock:
            create_backup()
            
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], EXCEL_FILE)
            file.save(filepath)
            with get_db_connection() as conn:
                cursor = conn.cursor()
                initialize_schema_from_excel(cursor)
            logging.info("Database re-initialized from uploaded file.")
        return jsonify({"message": "File uploaded and database re-initialized successfully."}), 200
    except Exception as e:
        logging.error(f"Error during file upload and re-initialization: {e}")
        return jsonify({"error": str(e)}), 500
    
def _generate_csv(db_df):
    output = io.BytesIO()
    db_df.to_csv(output, index=False)
    mimetype = 'text/csv'
    file_extension = 'csv'
    return output, mimetype, file_extension

def _generate_docx(db_df, file_name_from_req):
    output = io.BytesIO()
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    if file_name_from_req: document.add_heading(file_name_from_req, 0)
    table = document.add_table(rows=1, cols=len(db_df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(db_df.columns):
        hdr_cells[i].paragraphs[0].add_run(str(col_name)).bold = True
    for _, row in db_df.iterrows():
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = str(cell_value)
    document.save(output)
    mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    file_extension = 'docx'
    return output, mimetype, file_extension

def _generate_pdf(db_df, file_name_from_req):
    output = io.BytesIO()
    num_cols = len(db_df.columns)
    pagesize = landscape(A2) if num_cols > 25 else landscape(A3) if num_cols > 15 else landscape(A4)
    doc = SimpleDocTemplate(output, pagesize=pagesize, rightMargin=20, leftMargin=20, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    font_scale = max(0.5, 1 - (num_cols / 40.0))
    base_font_size = 8
    scaled_font_size = base_font_size * font_scale
    header_style = ParagraphStyle('Header', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=scaled_font_size, textColor=colors.white, alignment=TA_CENTER, leading=scaled_font_size * 1.2)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontName='Helvetica', fontSize=scaled_font_size - 1, alignment=TA_CENTER, leading=scaled_font_size * 1.2)
    elements = []
    available_width = doc.width
    col_widths = []
    for col in db_df.columns:
        header_len = len(str(col))
        max_data_len = db_df[col].astype(str).str.len().max()
        if pd.isna(max_data_len): max_data_len = 0
        col_width = max(header_len, int(max_data_len)) * scaled_font_size * 0.6
        col_widths.append(max(40, min(col_width, available_width / num_cols * 2.0 if num_cols > 0 else available_width)))
    total_width = sum(col_widths)
    if total_width > 0:
        col_widths = [w * available_width / total_width for w in col_widths]
    header_row = [Paragraph(str(h).replace('(₹)', '(INR)'), header_style) for h in db_df.columns]
    data_rows = [header_row]
    for _, row in db_df.iterrows():
        data_rows.append([Paragraph(str(item), body_style) for item in row])
    table = Table(data_rows, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.HexColor('#9B1C1C')),('TEXTCOLOR', (0,0), (-1,0), colors.black),('ALIGN', (0,0), (-1,-1), 'CENTER'),('VALIGN', (0,0), (-1,-1), 'MIDDLE'),('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),('BOTTOMPADDING', (0,0), (-1,0), 8),('TOPPADDING', (0,0), (-1,0), 8),('GRID', (0,0), (-1,-1), 0.5, colors.grey),('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#F0F0F0'), colors.white])]))
    elements.append(table)
    def header_footer(canvas, doc):
        canvas.saveState()
        if file_name_from_req:
            canvas.setFont('Helvetica-Bold', 12)
            canvas.drawCentredString(doc.width / 2.0 + doc.leftMargin, doc.height + doc.topMargin - 25, file_name_from_req)
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.grey)
        generation_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        canvas.drawString(doc.leftMargin, doc.bottomMargin - 15, f"Generated on: {generation_time}")
        canvas.drawRightString(doc.width + doc.leftMargin, doc.bottomMargin - 15, f"Page {doc.page}")
        canvas.restoreState()
    doc.build(elements, onFirstPage=header_footer, onLaterPages=header_footer)
    mimetype = 'application/pdf'
    file_extension = 'pdf'
    return output, mimetype, file_extension

def _generate_xlsx(db_df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        db_df.to_excel(writer, index=False, sheet_name='Contracts')
        workbook = writer.book
        worksheet = writer.sheets['Contracts']
        
        # --- THIS IS THE MODIFICATION ---
        # Create a bold header format
        header_format = workbook.add_format({'bold': True, 'fg_color': '#D7E4BC', 'border': 1, 'align': 'center', 'valign': 'vcenter'})
        # Create a center-aligned format for data cells
        cell_format = workbook.add_format({'align': 'center', 'valign': 'vcenter', 'border': 1})
        
        # Write headers and set column widths
        for col_num, value in enumerate(db_df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            column_len = max((db_df[value].astype(str).str.len().max() or 0), len(str(value))) + 2
            worksheet.set_column(col_num, col_num, column_len)
            
        # Write data rows with the center-aligned format
        for row_num, row_data in enumerate(db_df.values):
            for col_num, cell_data in enumerate(row_data):
                worksheet.write(row_num + 1, col_num, cell_data, cell_format)

        worksheet.freeze_panes(1, 0)
    mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    file_extension = 'xlsx'
    return output, mimetype, file_extension

@app.route('/api/export', methods=['GET'])
@token_required
def data_export():
    headers, field_types, _ = get_current_schema()
    if not headers:
        return jsonify({"error": "No schema defined, cannot export."}), 404

    params = request.args
    
    where_clauses, query_params = [], []
    filter_field = params.get('filterField')
    if filter_field:
        sanitized_filter_field = sanitize_column_name(filter_field)
        if filter_field in field_types.get('numeric', []):
            min_range, max_range = params.get('minRange'), params.get('maxRange')
            if min_range: where_clauses.append(f'CAST("{sanitized_filter_field}" AS REAL) >= ?'); query_params.append(float(min_range))
            if max_range: where_clauses.append(f'CAST("{sanitized_filter_field}" AS REAL) <= ?'); query_params.append(float(max_range))
        elif filter_field in field_types.get('date', []):
            from_date, to_date = params.get('fromDate'), params.get('toDate')
            if from_date: where_clauses.append(f'date("{sanitized_filter_field}") >= date(?)'); query_params.append(from_date)
            if to_date: where_clauses.append(f'date("{sanitized_filter_field}") <= date(?)'); query_params.append(to_date)
        else:
            filter_value = params.get('filterValue')
            if filter_value: where_clauses.append(f'LOWER("{sanitized_filter_field}") LIKE ?'); query_params.append(f"%{filter_value.lower()}%")
    
    where_statement = "WHERE " + " AND ".join(where_clauses) if where_clauses else ""

    sort_field = params.get('sortField', FIXED_START_COL)
    sort_direction = params.get('sortDirection', 'asc').upper()
    sanitized_sort_field = sanitize_column_name(sort_field)
    order_by_clause = f'ORDER BY CAST("{sanitize_column_name(FIXED_START_COL)}" AS REAL) ASC'
    if sort_field in field_types.get('numeric', []): order_by_clause = f'ORDER BY CAST("{sanitized_sort_field}" AS REAL) {sort_direction}'
    elif sort_field in field_types.get('date', []): order_by_clause = f'ORDER BY date("{sanitized_sort_field}") {sort_direction}'
    elif sort_field in field_types.get('text', []): order_by_clause = f'ORDER BY "{sanitized_sort_field}" {sort_direction}'

    with get_db_connection() as conn:
        sanitized_headers = [sanitize_column_name(h) for h in headers]
        select_clause = "SELECT " + ", ".join(f'"{col}"' for col in sanitized_headers)
        data_query = f"{select_clause} FROM {TABLE_NAME} {where_statement} {order_by_clause}"
        results = conn.execute(data_query, tuple(query_params)).fetchall()

    if not results:
        return jsonify({"error": "No data to export for the current filters."}), 404

    db_df = pd.DataFrame(results)
    db_df.columns = headers
    
    selected_fields_str = request.args.get('selectedFields', '')
    if selected_fields_str:
        selected_fields = selected_fields_str.split(',')
        # THIS IS THE FIX: Reorder the selected columns based on the original schema order.
        columns_to_keep = [header for header in headers if header in selected_fields]
        if columns_to_keep:
            db_df = db_df[columns_to_keep]

    format_type = request.args.get('format', 'xlsx')
    file_name_from_req = request.args.get('fileName')
    download_name_str = (file_name_from_req or 'contracts_export').replace(' ', '_')
    
    if format_type == 'csv':
        output, mimetype, file_extension = _generate_csv(db_df)
    elif format_type == 'docx':
        output, mimetype, file_extension = _generate_docx(db_df, file_name_from_req)
    elif format_type == 'pdf':
        output, mimetype, file_extension = _generate_pdf(db_df, file_name_from_req)
    else:
        output, mimetype, file_extension = _generate_xlsx(db_df)
    
    output.seek(0)
    return send_file(output, mimetype=mimetype, as_attachment=True, download_name=f'{download_name_str}.{file_extension}')

def validate_data(data_dict, field_types):
    errors = {}
    for key, value in data_dict.items():
        if value is None or str(value).strip() == '': continue
        if key in field_types.get('numeric', []):
            try: float(value)
            except (ValueError, TypeError): errors[key] = "must be a valid number."
        if key in field_types.get('date', []):
            try: pd.to_datetime(value, errors='raise')
            except (ValueError, TypeError): errors[key] = "must be a valid date format."
    if errors:
        return False, "Invalid data format: " + ", ".join([f"'{k}' ({v})" for k, v in errors.items()])
    return True, None

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    if not os.path.exists(BACKUP_FOLDER):
        os.makedirs(BACKUP_FOLDER)
    setup_database_and_schema()
    app.run(host='0.0.0.0', port=5001)